from docx import Document
from docx.shared import Inches

# Create document
doc = Document()
doc.add_heading("Test Document", 0)
doc.add_paragraph("This is a test paragraph for file2ai conversion testing.")
doc.add_paragraph("It includes multiple paragraphs to test text extraction.")

# Add a table
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = "Header 1"
table.cell(0, 1).text = "Header 2"
table.cell(1, 0).text = "Value 1"
table.cell(1, 1).text = "Value 2"

if __name__ == "__main__":
    import os

    # Ensure we're using absolute paths from project root
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    test_files_dir = os.path.join(project_root, "test_files")

    # Create test_files directory if it doesn't exist
    os.makedirs(test_files_dir, exist_ok=True)

    output_path = os.path.join(test_files_dir, "test.docx")
    doc.save(output_path)
    print(f"Test document created successfully at: {output_path}")
